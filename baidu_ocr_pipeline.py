"""
完整流水线：OCR课本 → 对照大纲 → 生成考点DOCX
"""
import os, sys, json, time, base64, io, re
from pathlib import Path

try:
    from PIL import Image
except ImportError:
    Image = None

try:
    import fitz
except ImportError:
    fitz = None

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

# ============ 配置 ============
API_KEY = "nQsyjGgNMeqDfbJ1VkcuDMrP"
SECRET_KEY = "DekBB85xIU5vvzfXZ1BoERYAtNwuSyh0"
TOKEN_URL = "https://aip.baidubce.com/oauth/2.0/token"
OCR_URL = "https://aip.baidubce.com/rest/2.0/ocr/v1/general_basic"

PDF_PATH = "F:/litellm/组织胚胎学/组织学与胚胎学 课本.pdf"
OUTLINE_PATH = "F:/litellm/组织胚胎学/大纲_text.txt"
OCR_OUTPUT = "F:/litellm/组织胚胎学/课本_ocr_full.txt"
DOCX_OUTPUT = "F:/litellm/组织胚胎学/组织学与胚胎学_考点整理.docx"
MODEL_DIR = "F:/litellm/.easyocr_models"
DEFAULT_DOCX = "F:/litellm/组织胚胎学/组织学与胚胎学_考点整理.docx"

# ============ 百度OCR API ============
def get_token():
    import urllib.request
    url = f"{TOKEN_URL}?grant_type=client_credentials&client_id={API_KEY}&client_secret={SECRET_KEY}"
    resp = urllib.request.urlopen(url, timeout=15)
    return json.loads(resp.read())["access_token"]

def ocr_image(img: Image.Image, token: str, retry=3) -> str:
    """调用百度OCR识别一张图片"""
    import urllib.request
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    img_base64 = base64.b64encode(buf.getvalue()).decode()

    data = urllib.parse.urlencode({"image": img_base64}).encode()
    req = urllib.request.Request(f"{OCR_URL}?access_token={token}", data=data)
    req.add_header("Content-Type", "application/x-www-form-urlencoded")

    for attempt in range(retry):
        try:
            resp = urllib.request.urlopen(req, timeout=30)
            result = json.loads(resp.read())
            if "words_result" in result:
                return "\n".join(item["words"] for item in result["words_result"])
            elif result.get("error_code") == 17:  # 每天免费额度用完了
                print(f"\n[!] 每日免费额度已用完! error: {result}")
                return None
            else:
                print(f"\n[!] API错误: {result}")
                return ""
        except Exception as e:
            if attempt < retry - 1:
                time.sleep(2 ** attempt)
                continue
            print(f"\n[!] 请求失败: {e}")
            return ""

# ============ OCR全本 ============
def ocr_entire_book(token, dpi=200, delay=0.8):
    """OCR整本课本"""
    if fitz is None:
        print("[错误] 需要pymupdf")
        sys.exit(1)

    doc = fitz.open(PDF_PATH)
    total = doc.page_count
    print(f"[课本] 共{total}页，开始OCR...")

    # 检查是否有已识别的结果，支持断点续传
    existing = {}
    if os.path.exists(OCR_OUTPUT):
        with open(OCR_OUTPUT, "r", encoding="utf-8") as f:
            content = f.read()
        for block in content.split("\n--- 第"):
            if block.strip():
                m = re.search(r"(\d+)页 ---", block)
                if m:
                    existing[int(m.group(1))] = block

    all_text = []
    success = 0
    fail = 0

    for i in range(total):
        page_num = i + 1
        if page_num in existing:
            all_text.append(f"--- 第{page_num}页 ---\n{existing[page_num].split('---', 1)[1].strip()}")
            continue

        page = doc[i]
        # 先检查是否有内置文字
        text = page.get_text().strip()
        if text and len(text) > 20:
            all_text.append(f"--- 第{page_num}页 ---\n{text}")
            print(f"  第{page_num}页 [文字] {len(text)}字")
            success += 1
            continue

        # OCR
        mat = fitz.Matrix(dpi / 72, dpi / 72)
        pix = page.get_pixmap(matrix=mat)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        print(f"  第{page_num}页 [OCR]", end="", flush=True)

        result = ocr_image(img, token)
        if result is None:  # 额度用完
            print(" -> 额度不足，暂停OCR")
            break
        if result:
            all_text.append(f"--- 第{page_num}页 ---\n{result}")
            print(f" {len(result)}字")
            success += 1
        else:
            print(" [空]")
            fail += 1

        # 进度
        if (page_num) % 20 == 0:
            print(f"\n[进度] {page_num}/{total}页 | 成功{success} | 空{fail}")
            # 阶段性保存
            with open(OCR_OUTPUT, "w", encoding="utf-8") as f:
                f.write("\n\n".join(all_text))

        time.sleep(delay)  # 避免超限

    doc.close()

    # 最终保存
    with open(OCR_OUTPUT, "w", encoding="utf-8") as f:
        f.write("\n\n".join(all_text))
    print(f"\n[完成] OCR完成: {success}页成功, {fail}页空")
    print(f"      输出: {OCR_OUTPUT}")

    # 读取后续继续
    return "\n\n".join(all_text) if all_text else ""


# ============ 提取大纲考点 ============
def parse_outline():
    """从大纲文本提取每章的考点"""
    if not os.path.exists(OUTLINE_PATH):
        print(f"[!] 大纲文件不存在: {OUTLINE_PATH}")
        return []

    with open(OUTLINE_PATH, "r", encoding="utf-8") as f:
        text = f.read()

    # 按章节拆分
    chapters = []
    # 匹配 "第X章 ..." 直到下一个"第X章"或末尾
    pattern = r"(第[\d一二三四五六七八九十]+章\s*.*?)(?=第[\d一二三四五六七八九十]+章\s|$)"
    matches = re.findall(pattern, text, re.DOTALL)

    for match in matches:
        lines = match.strip().split("\n")
        title = lines[0].strip() if lines else ""
        # 找出所有"掌握"和"熟悉"的内容
        master = []   # 掌握
        know = []     # 熟悉
        understand = []  # 了解
        content_points = []

        for line in lines:
            line = line.strip()
            if not line:
                continue
            # 提取教学要求
            if "掌握" in line and ("（" in line or "(" in line):
                # 处理多条目一行的情况
                parts = re.split(r"[（(]\d+[）)]", line)
                for p in parts:
                    p = p.strip().rstrip("。；；").strip()
                    # 拆开"掌握"和"熟悉"同时出现的行
                    if "。" in p:
                        sub_parts = p.split("。")
                        for s in sub_parts:
                            s = s.strip()
                            if "掌握" in s:
                                master.append(s.lstrip("掌握").lstrip("熟悉").strip())
                            elif "熟悉" in s:
                                know.append(s.replace("熟悉", "").strip())
                    else:
                        if "掌握" in p:
                            master.append(p.replace("掌握", "").replace("熟悉", "").strip().lstrip("，"))
                        elif "熟悉" in p:
                            know.append(p.replace("熟悉", "").replace("掌握", "").strip().lstrip("，"))

        chapters.append({
            "title": title,
            "master": master,
            "know": know,
            "full_text": match
        })

    return chapters


# ============ 用大纲对照课本整理 -> DOCX ============
def extract_chapter_text(ocr_text, chapter_title):
    """从OCR结果中提取对应章节的正文"""
    # 匹配章节标题
    patterns = [
        re.escape(chapter_title),
        chapter_title.replace(" ", r"\s*"),
    ]
    for p in patterns:
        m = re.search(p, ocr_text)
        if m:
            start = m.start()
            # 找下一个章节标题或结尾
            next_ch = re.search(r"--- 第\d+页 ---\s*第[\d一二三四五六七八九十]+章", ocr_text[start+len(chapter_title):])
            if next_ch:
                return ocr_text[start:start+len(chapter_title)+next_ch.start()]
            else:
                return ocr_text[start:start+min(5000, len(ocr_text)-start)]
    return ""


def organize_content(ocr_text, chapters):
    """按章节整理知识点"""
    if not ocr_text or not chapters:
        return chapters

    for ch in chapters:
        # 提取本章的正文
        ch_text = extract_chapter_text(ocr_text, ch["title"])
        if ch_text:
            ch["ocr_excerpt"] = ch_text
        else:
            ch["ocr_excerpt"] = ""

    return chapters


def generate_docx(chapters, output_path):
    """生成考点整理DOCX"""
    if Document is None:
        print("[!] python-docx未安装: pip install python-docx")
        print("  改用纯文本输出...")
        _generate_txt(chapters)
        return

    doc = Document()

    # 标题
    title = doc.add_heading("组织学与胚胎学 考点整理", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("根据教学大纲要求整理")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    chapter_count = 0
    for ch in chapters:
        if not ch["master"] and not ch["know"]:
            continue
        chapter_count += 1

        # 章节标题
        doc.add_heading(ch["title"], level=1)

        # (一) 掌握内容
        if ch["master"]:
            h = doc.add_heading("【掌握】", level=2)
            for item in ch["master"]:
                if len(item) > 3:
                    p = doc.add_paragraph(item, style="List Bullet")

        # (二) 熟悉内容
        if ch["know"]:
            h = doc.add_heading("【熟悉】", level=2)
            for item in ch["know"]:
                if len(item) > 3:
                    p = doc.add_paragraph(item, style="List Bullet")

        # (三) 考点细节（如果有OCR内容）
        if ch.get("ocr_excerpt"):
            h = doc.add_heading("【知识点详解】", level=2)
            # 取关键段落
            text = ch["ocr_excerpt"]
            # 去掉章节标题行
            text = re.sub(r"--- 第\d+页 ---", "", text)
            text = re.sub(r"第[\d一二三四五六七八九十]+章.*", "", text)
            # 按句号分段
            sentences = re.split(r"[。！？]", text)
            for sent in sentences:
                sent = sent.strip()
                if len(sent) > 10:
                    doc.add_paragraph(sent + "。", style="List Bullet")

        doc.add_page_break()

    doc.save(output_path)
    print(f"\n[DOCX] 已生成: {output_path}")
    print(f"       共{chapter_count}章")


def _generate_txt(chapters, output="F:/litellm/组织胚胎学/考点整理.txt"):
    """备用：生成纯文本"""
    lines = []
    lines.append("=" * 50)
    lines.append("组织学与胚胎学 考点整理")
    lines.append("=" * 50)
    lines.append("")

    for ch in chapters:
        if not ch["master"] and not ch["know"]:
            continue
        lines.append(f"\n{'='*40}")
        lines.append(ch["title"])
        lines.append(f"{'='*40}")

        if ch["master"]:
            lines.append("\n【掌握】")
            for item in ch["master"]:
                if len(item) > 3:
                    lines.append(f"  • {item}")

        if ch["know"]:
            lines.append("\n【熟悉】")
            for item in ch["know"]:
                if len(item) > 3:
                    lines.append(f"  • {item}")

        if ch.get("ocr_excerpt"):
            lines.append("\n【知识点】")
            text = re.sub(r"--- 第\d+页 ---", "", ch["ocr_excerpt"])
            sentences = re.split(r"[。！？]", text)
            for sent in sentences:
                sent = sent.strip()
                if len(sent) > 10:
                    lines.append(f"  - {sent}。")

    with open(output, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"\n[TXT] 已生成: {output}")


# ============ 主流程 ============
def main():
    print("=" * 50)
    print("组织学与胚胎学 考点整理流水线")
    print("=" * 50)

    # Step 1: 获取Token
    print("\n[1/4] 获取百度OCR Token...")
    token = get_token()
    print(f"  Token获取成功 ({token[:20]}...)")

    # Step 2: OCR全本
    print("\n[2/4] OCR课本全文...")
    ocr_text = ocr_entire_book(token, dpi=200, delay=1.0)

    if not ocr_text or len(ocr_text) < 100:
        print("[!] OCR未能获取足够文字，检查API额度")
        # 尝试读取已保存的
        if os.path.exists(OCR_OUTPUT):
            with open(OCR_OUTPUT, "r", encoding="utf-8") as f:
                ocr_text = f.read()
            print(f"  读取已有OCR结果: {len(ocr_text)}字")

    # Step 3: 解析大纲考点
    print("\n[3/4] 提取大纲考点...")
    chapters = parse_outline()
    if not chapters:
        print("[!] 大纲解析失败")
        sys.exit(1)
    print(f"  解析到{len(chapters)}章")

    # 展示各章考点
    for ch in chapters:
        m = len(ch["master"]) if ch["master"] else 0
        k = len(ch["know"]) if ch["know"] else 0
        if m > 0 or k > 0:
            print(f"  {ch['title']}: 掌握{m}条, 熟悉{k}条")

    # Step 4: 生成DOCX
    print("\n[4/4] 生成考点文档...")
    chapters = organize_content(ocr_text, chapters)
    generate_docx(chapters, DOCX_OUTPUT)


if __name__ == "__main__":
    main()
